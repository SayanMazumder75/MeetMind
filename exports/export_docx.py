"""
export_docx.py — Export meeting notes to a Word document (.docx).

Setup:
    pip install python-docx
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from core.config import config
from core.transcriber import TranscriptChunk
from core.summarizer import MeetingSummary

logger = logging.getLogger(__name__)


def export_to_docx(
    session_id: str,
    transcript: list[TranscriptChunk],
    summary: Optional[MeetingSummary],
    output_path: Optional[Path] = None,
    meeting_title: str = "Meeting Notes",
) -> Path:
    """
    Generate a .docx Word document with the meeting notes.
    Returns the path to the created file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    if output_path is None:
        output_path = config.SUMMARY_DIR / f"meeting_notes_{session_id}.docx"

    doc = Document()

    # ── Styles ─────────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ──────────────────────────────────────────────────────────────────
    title_para = doc.add_heading(meeting_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(
        f"Session: {session_id}   |   Generated: {datetime.now().strftime('%d %b %Y %H:%M')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = RGBColor(120, 120, 120)
    meta.runs[0].font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── Summary ────────────────────────────────────────────────────────────────
    if summary:
        doc.add_heading("Meeting Summary", level=1)
        doc.add_paragraph(summary.summary)

        if summary.action_items:
            doc.add_heading("Action Items", level=2)
            for item in summary.action_items:
                p = doc.add_paragraph(item, style="List Bullet")
                p.runs[0].font.size = Pt(11)

        if summary.key_points:
            doc.add_heading("Key Points", level=2)
            for point in summary.key_points:
                doc.add_paragraph(point, style="List Bullet")

        if summary.highlights:
            doc.add_heading("Highlights", level=2)
            for h in summary.highlights:
                p = doc.add_paragraph()
                run = p.add_run(f'"{h}"')
                run.font.italic = True
                run.font.color.rgb = RGBColor(60, 60, 120)

    # ── Transcript ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Full Transcript", level=1)

    for chunk in transcript:
        p = doc.add_paragraph()
        # Timestamp
        ts_run = p.add_run(f"[{chunk.time_label}] ")
        ts_run.font.color.rgb = RGBColor(100, 100, 180)
        ts_run.font.size = Pt(9)
        # Speaker
        sp_run = p.add_run(f"{chunk.speaker}: ")
        sp_run.bold = True
        sp_run.font.size = Pt(10)
        # Text
        txt_run = p.add_run(chunk.text)
        txt_run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    doc.save(str(output_path))
    logger.info(f"DOCX exported: {output_path}")
    return output_path
